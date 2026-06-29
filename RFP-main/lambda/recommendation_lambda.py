"""
Strands Agents - Tool 6: Recommendation Lambda
Step 7: Return top 2 suppliers, generate .docx recommendation report,
        set approval_required flag, auto-award or set pending_approval in DynamoDB
"""

import json
import logging
import boto3
import os
import io
from datetime import datetime
from typing import List, Dict, Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger()
logger.setLevel(logging.INFO)

REGION             = os.environ.get("REGION", "us-east-1")
RFP_DOCS_BUCKET    = os.environ.get("RFP_DOCS_BUCKET", "rfp-documents-quadrasystems")
REQUESTS_TABLE     = os.environ.get("REQUESTS_TABLE", "rfp-requests")

s3         = boto3.client("s3", region_name=REGION)
dynamodb   = boto3.resource("dynamodb", region_name=REGION)


def handler(event, context):
    """
    Recommendation Lambda Handler

    Input: {"body": {"rfp_id": "RFP-...", "scored_proposals": [...]}}
    Output: Top 2 recommendations + .docx report presigned URL + approval decision
    """
    try:
        logger.info("[Tool 6] 🏆 Recommendation - Starting")

        body = event.get("body", {})
        if isinstance(body, str):
            body = json.loads(body)

        rfp_id           = body.get("rfp_id", "")
        scored_proposals = body.get("scored_proposals", [])

        if not rfp_id or not scored_proposals:
            logger.error("[Tool 6] ❌ Missing rfp_id or scored_proposals")
            return _response(400, {
                "error": "rfp_id and scored_proposals required",
                "success": False
            })

        logger.info(f"[Tool 6] Evaluating {len(scored_proposals)} proposals")

        if len(scored_proposals) < 1:
            return _response(200, {
                "success": True,
                "recommendation_count": 0,
                "recommendations": [],
                "approval_required": False,
                "timestamp": datetime.utcnow().isoformat()
            })

        # Sort by total_score descending — highest score = Rank 1
        scored_proposals = sorted(
            scored_proposals,
            key=lambda x: float(x.get("total_score", 0)),
            reverse=True
        )

        # Top 2
        top_2 = scored_proposals[:2]
        recommendations: List[Dict[str, Any]] = []
        all_flags: List[str] = []

        for rank, proposal in enumerate(top_2, 1):
            flags = proposal.get("risk_flags", [])
            all_flags.extend(flags)

            rec = {
                "rank":         rank,
                "supplier_id":  proposal.get("supplier_id", ""),
                "supplier_name": proposal.get("supplier_name", proposal.get("supplier_id", "")),
                "total_score":  proposal.get("total_score", 0),
                "price":        proposal.get("price", 0),
                "delivery_days": proposal.get("delivery_days", 0),
                "quality":      proposal.get("quality", 0),
                "compliance":   proposal.get("compliance", "No"),
                "price_score":  proposal.get("score_breakdown", {}).get("price_score", 0),
                "quality_score": proposal.get("score_breakdown", {}).get("quality_score", 0),
                "delivery_score": proposal.get("score_breakdown", {}).get("delivery_score", 0),
                "compliance_score": proposal.get("score_breakdown", {}).get("compliance_score", 0),
                "risk_flags":   flags,
                "flags":        flags,
                "reasoning":    _generate_reasoning(proposal, rank),
                "recommendation": _get_recommendation_text(rank, flags)
            }
            recommendations.append(rec)
            logger.info(f"[Tool 6] Rank {rank}: {rec['supplier_id']} score={rec['total_score']}")

        # Approval required if any risk flags exist
        approval_required = len(all_flags) > 0

        # Generate .docx recommendation report
        rec_url = ""
        try:
            buffer = _generate_recommendation_docx(
                rfp_id=rfp_id,
                top2=recommendations,
                all_flags=list(set(all_flags)),
                approval_required=approval_required
            )
            rec_key = f"recommendations/{rfp_id}-recommendation.docx"
            s3.put_object(
                Bucket=RFP_DOCS_BUCKET,
                Key=rec_key,
                Body=buffer.getvalue(),
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            rec_url = s3.generate_presigned_url(
                "get_object",
                Params={"Bucket": RFP_DOCS_BUCKET, "Key": rec_key},
                ExpiresIn=604800  # 7 days
            )
            logger.info(f"[Tool 6] 💾 Recommendation .docx saved: {rec_key}")
        except Exception as e:
            logger.warning(f"[Tool 6] ⚠️ .docx generation failed: {e}")

        # FIX 7: Update DynamoDB rfp-requests with award or pending status
        _update_request_status(rfp_id, approval_required, top_2)

        result = {
            "status":           "success",
            "top_2":            [
                {
                    "rank":          rec["rank"],
                    "supplier_id":   rec["supplier_id"],
                    "supplier_name": rec["supplier_name"],
                    "total_score":   rec["total_score"],
                    "flags":         rec["flags"]
                }
                for rec in recommendations
            ],
            "approval_required": approval_required,
            "rec_docx_url":      rec_url
        }

        logger.info(
            f"[Tool 6] ✅ {len(recommendations)} recommendations | "
            f"approval_required={approval_required}"
        )
        return _response(200, result)

    except Exception as e:
        logger.error(f"[Tool 6] ❌ Error: {e}", exc_info=True)
        return _response(500, {"error": str(e), "success": False})


# ============================================================================
# DOCX REPORT GENERATOR
# ============================================================================

def _generate_recommendation_docx(
    rfp_id: str,
    top2: List[Dict[str, Any]],
    all_flags: List[str],
    approval_required: bool
) -> io.BytesIO:
    """Generate a formatted .docx recommendation report."""
    doc = Document()

    # Title
    title = doc.add_heading("SUPPLIER RECOMMENDATION REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(f"RFP ID: {rfp_id}", level=1)
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(
        f"Approval Required: {'YES — Human sign-off needed' if approval_required else 'NO — Auto-approved'}"
    )

    doc.add_heading("Top Supplier Recommendations", level=1)

    for s in top2:
        doc.add_heading(f"Rank #{s['rank']}: {s['supplier_name'] or s['supplier_id']}", level=2)

        t = doc.add_table(rows=6, cols=2)
        t.style = "Table Grid"
        rows = [
            ("Total Score",      f"{s['total_score']}/100"),
            ("Price Score",      str(s.get("price_score", ""))),
            ("Quality Score",    str(s.get("quality_score", ""))),
            ("Delivery Score",   str(s.get("delivery_score", ""))),
            ("Compliance Score", str(s.get("compliance_score", ""))),
            ("Risk Flags",       ", ".join(s["flags"]) if s["flags"] else "None")
        ]
        for i, (k, v) in enumerate(rows):
            t.rows[i].cells[0].text = k
            t.rows[i].cells[1].text = str(v)

        doc.add_paragraph(f"Reasoning: {s['reasoning']}")
        doc.add_paragraph(f"Recommendation: {s['recommendation']}")
        doc.add_paragraph()

    # Risk flags section
    if all_flags:
        doc.add_heading("Risk Flags Detected", level=2)
        for f in all_flags:
            doc.add_paragraph(f"• {f}")
        doc.add_paragraph()

    # Approval decision
    doc.add_heading("Approval Decision", level=2)
    doc.add_paragraph(
        "PENDING HUMAN APPROVAL — Please review risk flags before proceeding."
        if approval_required
        else "AUTO-APPROVED — No critical risks detected. Contract can proceed."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================================
# CONTRACT AWARD — FIX 7
# ============================================================================

def _update_request_status(
    rfp_id: str,
    approval_required: bool,
    top_2: List[Dict[str, Any]]
) -> None:
    """
    FIX 7: Update DynamoDB rfp-requests with award or pending_approval status.
    - approval_required = False → status = 'awarded', record winning supplier
    - approval_required = True  → status = 'pending_approval'
    """
    try:
        table = dynamodb.Table(REQUESTS_TABLE)
        now = datetime.utcnow().isoformat()

        if not approval_required and top_2:
            winner = top_2[0]
            table.update_item(
                Key={"RequestID": rfp_id},
                UpdateExpression=(
                    "SET #s = :s, awarded_supplier = :sup, "
                    "awarded_supplier_score = :score, awarded_at = :at"
                ),
                ExpressionAttributeNames={"#s": "Status"},
                ExpressionAttributeValues={
                    ":s":     "awarded",
                    ":sup":   winner.get("supplier_id", "unknown"),
                    ":score": str(winner.get("total_score", 0)),
                    ":at":    now
                }
            )
            logger.info(
                f"[Tool 6] 🏅 Contract AUTO-AWARDED to "
                f"{winner.get('supplier_id')} (score={winner.get('total_score')})"
            )
        else:
            table.update_item(
                Key={"RequestID": rfp_id},
                UpdateExpression="SET #s = :s, pending_since = :at",
                ExpressionAttributeNames={"#s": "Status"},
                ExpressionAttributeValues={
                    ":s":  "pending_approval",
                    ":at": now
                }
            )
            logger.info(f"[Tool 6] ⏳ Status set to pending_approval for {rfp_id}")

    except Exception as e:
        logger.warning(f"[Tool 6] ⚠️ DynamoDB status update failed: {e}")


# ============================================================================
# HELPERS
# ============================================================================

def _generate_reasoning(proposal: Dict[str, Any], rank: int) -> str:
    score = proposal.get("total_score", 0)
    risk_count = len(proposal.get("risk_flags", []))
    risk_note = f" ({risk_count} risk flag(s) detected)" if risk_count else " (no risk flags)"
    return (
        f"Rank {rank} supplier scores {score:.1f}/100 based on weighted criteria "
        f"(Price 30%, Quality 30%, Delivery 20%, Compliance 20%){risk_note}"
    )


def _get_recommendation_text(rank: int, risk_flags: List[str]) -> str:
    if rank == 1:
        return (
            f"RECOMMENDED with risk review — {len(risk_flags)} flag(s)"
            if risk_flags else "RECOMMENDED — Optimal choice"
        )
    return (
        f"BACKUP option — Review {len(risk_flags)} risk(s) before proceeding"
        if risk_flags else "BACKUP option — Good alternative"
    )


def _response(status_code: int, body: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "statusCode": status_code,
        "headers": {"Content-Type": "application/json"},
        "body": json.dumps(body, default=str)
    }
