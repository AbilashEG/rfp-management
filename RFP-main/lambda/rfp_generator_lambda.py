"""
Strands Agents - Tool 2: RFP Generator Lambda
Step 3: Generate .docx RFP document → S3 → DynamoDB
Returns presigned URL for the .docx file
"""

import json
import logging
import boto3
import os
import io
import uuid
from datetime import datetime
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger()
logger.setLevel(logging.INFO)

REGION = os.environ.get("REGION", "us-east-1")
RFP_DOCS_BUCKET = os.environ.get("RFP_DOCS_BUCKET", "rfp-documents-quadrasystems")
RFP_DOCS_PREFIX = "rfp-documents/"

dynamodb = boto3.resource("dynamodb", region_name=REGION)
s3 = boto3.client("s3", region_name=REGION)
requests_table = dynamodb.Table("rfp-requests")


def handler(event, context):
    """
    RFP Generator Lambda Handler

    Input: {"body": {"rfp_id": "RFP-...", "requirement": "...", "supplier_ids": [...]}}
    Output: Generated .docx RFP with S3 presigned URL
    """
    try:
        logger.info("[Tool 2] 📄 RFP Generator - Starting")

        body = event.get("body", {})
        if isinstance(body, str):
            body = json.loads(body)

        requirement = body.get("requirement", "").strip()
        supplier_ids = body.get("supplier_ids", [])
        rfp_id = body.get("rfp_id", f"RFP-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}")
        component_name = body.get("component_name", requirement[:60] if requirement else "Component")
        specs = body.get("specs", "As per requirement")
        quantity = body.get("quantity", "As specified")
        deadline = body.get("deadline", "2026-09-30")
        google_form_url = body.get("google_form_url", "")

        if not requirement or not supplier_ids:
            return _response(400, {"error": "requirement and supplier_ids required", "success": False})

        logger.info(f"[Tool 2] Generating RFP: {rfp_id} for {len(supplier_ids)} suppliers")

        timestamp = datetime.now().isoformat()

        # Generate .docx
        buffer = _generate_rfp_docx(
            rfp_id=rfp_id,
            component_name=component_name,
            specs=specs,
            quantity=str(quantity),
            deadline=deadline,
            supplier_ids=supplier_ids,
            requirement=requirement,
            google_form_url=google_form_url
        )

        # Save .docx to S3
        s3_key = f"{RFP_DOCS_PREFIX}{rfp_id}.docx"
        presigned_url = ""
        try:
            s3.put_object(
                Bucket=RFP_DOCS_BUCKET,
                Key=s3_key,
                Body=buffer.getvalue(),
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            logger.info(f"[Tool 2] 💾 Saved .docx to S3: s3://{RFP_DOCS_BUCKET}/{s3_key}")

            presigned_url = s3.generate_presigned_url(
                "get_object",
                Params={"Bucket": RFP_DOCS_BUCKET, "Key": s3_key},
                ExpiresIn=604800  # 7 days
            )
            logger.info(f"[Tool 2] 🔗 Presigned URL generated (7 days)")
        except Exception as e:
            logger.warning(f"[Tool 2] ⚠️ S3 save failed: {e}")

        # Save metadata to DynamoDB
        try:
            requests_table.put_item(
                Item={
                    "RequestID": rfp_id,
                    "CreatedAt": timestamp,
                    "Requirement": requirement,
                    "SupplierCount": len(supplier_ids),
                    "Status": "Active",
                    "S3Location": s3_key,
                    "DocxPresignedUrl": presigned_url,
                    "GoogleFormUrl": google_form_url
                }
            )
            logger.info("[Tool 2] 💾 Saved metadata to DynamoDB rfp-requests")
        except Exception as e:
            logger.warning(f"[Tool 2] ⚠️ DynamoDB save failed: {e}")

        result = {
            "status": "success",
            "rfp_id": rfp_id,
            "docx_url": presigned_url
        }

        logger.info(f"[Tool 2] ✅ RFP .docx generated: {rfp_id}")
        return _response(200, result)

    except Exception as e:
        logger.error(f"[Tool 2] ❌ Error: {e}", exc_info=True)
        return _response(500, {"error": str(e), "success": False})


def _generate_rfp_docx(
    rfp_id: str,
    component_name: str,
    specs: str,
    quantity: str,
    deadline: str,
    supplier_ids: List[str],
    requirement: str,
    google_form_url: str
) -> io.BytesIO:
    """Generate a professional .docx RFP document matching real-world RFP standards."""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Helper: add horizontal rule ───────────────────────────
    def add_hr(para):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C4B89E')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Styles ─────────────────────────────────────────────────
    DARK   = RGBColor(0x2C, 0x2C, 0x2C)
    ORANGE = RGBColor(0xE8, 0xA0, 0x20)
    MUTED  = RGBColor(0x7A, 0x72, 0x65)

    def set_font(run, size, bold=False, color=None, italic=False):
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

    issued_date = datetime.utcnow().strftime('%B %d, %Y')

    # ══════════════════════════════════════════════════════════
    # HEADER — Organization + RFP label
    # ══════════════════════════════════════════════════════════
    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = org_para.add_run('QUADRA SYSTEMS')
    set_font(r, 18, bold=True, color=DARK)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_para.add_run('Procurement & Supply Chain Division')
    set_font(r, 10, color=MUTED)

    doc.add_paragraph()
    add_hr(doc.add_paragraph())

    # ══════════════════════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('REQUEST FOR PROPOSAL')
    set_font(r, 22, bold=True, color=DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(component_name.upper())
    set_font(r, 13, bold=False, color=ORANGE)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # REFERENCE BOX TABLE
    # ══════════════════════════════════════════════════════════
    ref_table = doc.add_table(rows=4, cols=2)
    ref_table.style = 'Table Grid'
    ref_data = [
        ('RFP Reference Number', rfp_id),
        ('Issue Date',           issued_date),
        ('Submission Deadline',  f'{deadline} at 5:00 PM IST'),
        ('Number of Suppliers',  str(len(supplier_ids))),
    ]
    for i, (k, v) in enumerate(ref_data):
        ref_table.rows[i].cells[0].text = k
        ref_table.rows[i].cells[1].text = v
        for cell in ref_table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)

    doc.add_paragraph()
    add_hr(doc.add_paragraph())

    # ══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════
    toc_heading = doc.add_paragraph()
    r = toc_heading.add_run('TABLE OF CONTENTS')
    set_font(r, 11, bold=True, color=DARK)

    toc_items = [
        '1.  Procurement Overview',
        '2.  Requirement Specifications',
        '3.  Scope of Supply',
        '4.  Supplier Eligibility Criteria',
        '5.  Evaluation Criteria & Weightage',
        '6.  Proposal Submission Requirements',
        '7.  Terms and Conditions',
        '8.  Submission Instructions',
        '9.  Contact Information',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = MUTED

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 1 — PROCUREMENT OVERVIEW
    # ══════════════════════════════════════════════════════════
    def section_heading(text, num):
        p = doc.add_paragraph()
        r = p.add_run(f'{num}.  {text}')
        set_font(r, 13, bold=True, color=DARK)
        add_hr(doc.add_paragraph())

    def body_text(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.2)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
        return p

    section_heading('PROCUREMENT OVERVIEW', 1)
    body_text(
        'Quadra Systems ("the Company") invites proposals from qualified suppliers for the '
        f'supply of {component_name}. This Request for Proposal (RFP) outlines the technical '
        'specifications, evaluation criteria, and terms under which proposals will be considered. '
        'The Company is committed to selecting suppliers based on quality, price competitiveness, '
        'delivery reliability, and regulatory compliance.'
    )
    doc.add_paragraph()
    body_text(
        f'This procurement is initiated in response to the following requirement: {requirement}'
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 2 — REQUIREMENT SPECIFICATIONS
    # ══════════════════════════════════════════════════════════
    section_heading('REQUIREMENT SPECIFICATIONS', 2)

    spec_table = doc.add_table(rows=6, cols=2)
    spec_table.style = 'Table Grid'
    spec_data = [
        ('Component / Item',         component_name),
        ('Technical Specifications', specs),
        ('Required Quantity',        f'{quantity} units'),
        ('Delivery Deadline',        deadline),
        ('Delivery Location',        'Quadra Systems, India'),
        ('Packing Requirements',     'As per IATF 16949 / ISO 9001 standards'),
    ]
    for i, (k, v) in enumerate(spec_data):
        cell_k = spec_table.rows[i].cells[0]
        cell_v = spec_table.rows[i].cells[1]
        cell_k.text = k
        cell_v.text = v
        for run in cell_k.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        for run in cell_v.paragraphs[0].runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 3 — SCOPE OF SUPPLY
    # ══════════════════════════════════════════════════════════
    section_heading('SCOPE OF SUPPLY', 3)
    body_text('The selected supplier shall be responsible for the following:')
    doc.add_paragraph()

    scope_items = [
        f'Supply of {component_name} as per specifications outlined in Section 2.',
        'Provision of quality certificates, test reports, and compliance documentation.',
        'Packaging, labelling, and shipping as per Quadra Systems packaging standards.',
        'Delivery within the agreed timeline with no delays. Late delivery penalties may apply.',
        'After-delivery support and warranty as per the mutually agreed terms.',
        'Compliance with all applicable environmental, safety, and industry regulations.',
    ]
    for item in scope_items:
        p = doc.add_paragraph(f'•  {item}')
        p.paragraph_format.left_indent = Inches(0.4)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 4 — SUPPLIER ELIGIBILITY
    # ══════════════════════════════════════════════════════════
    section_heading('SUPPLIER ELIGIBILITY CRITERIA', 4)
    body_text('To be considered for this RFP, suppliers must meet the following minimum requirements:')
    doc.add_paragraph()

    eligibility = [
        'Valid business registration and GST / tax compliance documentation.',
        'Minimum 3 years of experience supplying components to the automotive or manufacturing sector.',
        'ISO 9001:2015 certification (IATF 16949 preferred).',
        'Demonstrated capacity to fulfill the required quantity within the specified timeline.',
        'No history of regulatory violations or supply chain fraud.',
        'Willingness to participate in Quadra Systems supplier audit process.',
    ]
    for item in eligibility:
        p = doc.add_paragraph(f'•  {item}')
        p.paragraph_format.left_indent = Inches(0.4)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 5 — EVALUATION CRITERIA
    # ══════════════════════════════════════════════════════════
    section_heading('EVALUATION CRITERIA & WEIGHTAGE', 5)
    body_text(
        'Proposals will be evaluated using the following weighted scoring model. '
        'The supplier with the highest total weighted score will be recommended for award.'
    )
    doc.add_paragraph()

    eval_table = doc.add_table(rows=6, cols=3)
    eval_table.style = 'Table Grid'
    eval_headers = ['Criteria', 'Weightage', 'Scoring Basis']
    for j, h in enumerate(eval_headers):
        eval_table.rows[0].cells[j].text = h
        for run in eval_table.rows[0].cells[j].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    eval_data = [
        ('Unit Price',          '30%', 'Lowest price receives highest score'),
        ('Quality Score',       '30%', 'Based on certifications and past performance'),
        ('Delivery Time',       '20%', 'Shorter lead time receives higher score'),
        ('Regulatory Compliance','20%','ISO, IATF, and other certifications verified'),
        ('TOTAL',               '100%',''),
    ]
    for i, (c, w, b) in enumerate(eval_data, 1):
        eval_table.rows[i].cells[0].text = c
        eval_table.rows[i].cells[1].text = w
        eval_table.rows[i].cells[2].text = b
        for cell in eval_table.rows[i].cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 6 — PROPOSAL REQUIREMENTS
    # ══════════════════════════════════════════════════════════
    section_heading('PROPOSAL SUBMISSION REQUIREMENTS', 6)
    body_text('Each supplier must submit a proposal including the following information:')
    doc.add_paragraph()

    proposal_req = [
        'Company profile and registration documents.',
        'Technical compliance statement against specifications in Section 2.',
        'Unit price (exclusive of taxes), with breakup of cost components.',
        'Confirmed lead time and delivery commitment.',
        'Quality certifications (ISO 9001, IATF 16949, or equivalent).',
        'Three client references from similar supply engagements.',
        'Sample product or datasheet (if available).',
        'Any deviations or exceptions to this RFP must be clearly stated.',
    ]
    for item in proposal_req:
        p = doc.add_paragraph(f'•  {item}')
        p.paragraph_format.left_indent = Inches(0.4)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 7 — TERMS AND CONDITIONS
    # ══════════════════════════════════════════════════════════
    section_heading('TERMS AND CONDITIONS', 7)

    terms = [
        ('Confidentiality',
         'All information contained in this RFP is confidential and shall not be disclosed to third parties.'),
        ('No Commitment',
         'This RFP does not constitute a commitment by Quadra Systems to award a contract. '
         'The Company reserves the right to accept or reject any proposal.'),
        ('Ownership',
         'All proposals submitted become the property of Quadra Systems.'),
        ('False Information',
         'Any supplier found to have submitted false or misleading information will be immediately disqualified.'),
        ('Payment Terms',
         'Standard payment terms are Net 30 days from delivery and acceptance of goods unless otherwise agreed.'),
        ('Warranty',
         'Supplier shall warrant that all goods supplied are free from defects in materials and workmanship '
         'for a period of 12 months from date of delivery.'),
    ]
    for term, desc in terms:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f'{term}:  ')
        set_font(r1, 10.5, bold=True, color=DARK)
        r2 = p.add_run(desc)
        set_font(r2, 10.5, color=DARK)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 8 — SUBMISSION INSTRUCTIONS
    # ══════════════════════════════════════════════════════════
    section_heading('SUBMISSION INSTRUCTIONS', 8)

    sub_instructions = [
        f'Proposals must be submitted no later than {deadline} at 5:00 PM IST.',
        'Proposals must be submitted via the Google Form link provided in the accompanying email.',
        'Include all required documents as attachments in the submission form.',
        'Late submissions will not be accepted under any circumstances.',
        'Shortlisted suppliers will be notified within 7 business days of the submission deadline.',
        'Quadra Systems may request a clarification meeting or additional documentation after review.',
    ]
    for item in sub_instructions:
        p = doc.add_paragraph(f'•  {item}')
        p.paragraph_format.left_indent = Inches(0.4)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)

    doc.add_paragraph()
    if google_form_url:
        gf_para = doc.add_paragraph()
        gf_para.paragraph_format.left_indent = Inches(0.2)
        r = gf_para.add_run('Proposal Submission Link:  ')
        set_font(r, 10.5, bold=True, color=ORANGE)
        r2 = gf_para.add_run(google_form_url)
        set_font(r2, 10.5, color=RGBColor(0x00, 0x56, 0xD2))

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 9 — CONTACT INFORMATION
    # ══════════════════════════════════════════════════════════
    section_heading('CONTACT INFORMATION', 9)
    body_text(
        'For clarifications regarding technical specifications, please contact the Procurement Team. '
        'All questions must be submitted in writing and will be responded to within 2 business days.'
    )
    doc.add_paragraph()

    contact_table = doc.add_table(rows=4, cols=2)
    contact_table.style = 'Table Grid'
    contact_data = [
        ('Organization',   'Quadra Systems'),
        ('Department',     'Procurement & Supply Chain'),
        ('Email',          'procurement@quadrasystems.com'),
        ('RFP Reference',  rfp_id),
    ]
    for i, (k, v) in enumerate(contact_data):
        contact_table.rows[i].cells[0].text = k
        contact_table.rows[i].cells[1].text = v
        for run in contact_table.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        for run in contact_table.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    doc.add_paragraph()
    add_hr(doc.add_paragraph())

    # ── Footer note ────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_para.add_run(
        f'This document is confidential and intended solely for invited suppliers.  '
        f'RFP Reference: {rfp_id}  ·  Issued: {issued_date}'
    )
    set_font(r, 9, italic=True, color=MUTED)

    # ── Save ───────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _response(status_code: int, body: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "statusCode": status_code,
        "headers": {"Content-Type": "application/json"},
        "body": json.dumps(body, default=str)
    }
